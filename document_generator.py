from docx import Document
import os
from datetime import datetime
import re


def clean_filename(name):
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = name.replace(" ", "_")
    return name[:50]


def create_word_document(title, content):
    os.makedirs("output", exist_ok=True)

    safe_title = clean_filename(title)

    file_name = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join("output", file_name)

    doc = Document()
    doc.add_heading(title.replace('"', ''), 0)

    for section in content.split("\n"):
        if section.strip():
            doc.add_paragraph(section.strip())

    doc.save(file_path)

    return file_path